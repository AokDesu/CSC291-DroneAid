"""Build docs/03-personas.docx — two user personas for DroneAid."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        if level == 0:
            run.font.size = Pt(26)
        elif level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    return p


def add_para(doc, text: str, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items: list[str]):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.runs[0] if p.runs else p.add_run(it)
        if not p.runs:
            p.add_run(it)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
        if not p.runs[0].text:
            p.runs[0].text = it


def add_kv_table(doc, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.width = Inches(1.7)
        c2.width = Inches(4.8)
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.name = "Arial"
        r2.font.size = Pt(11)
        shade_cell(c1, "EAF3FB")
        set_cell_border(c1)
        set_cell_border(c2)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


def add_quote_box(doc, quote: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"“{quote}”")
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    shade_cell(cell, "EAF3FB")
    set_cell_border(cell, color="A6C8E0")


def add_photo_placeholder(doc, label: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Inches(2.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ Photo placeholder ]\n{label}")
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    shade_cell(cell, "F2F2F2")
    set_cell_border(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build():
    out = Path(__file__).resolve().parent.parent / "docs" / "03-personas.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Title page
    add_heading(doc, "DroneAid — User Personas", level=0)
    add_para(doc, "CSC291 group project, KMUTT", italic=True, size=11)
    add_para(doc, "Team: Aok (Lead), Belle, Bew, Poom, Tawan", italic=True, size=11)
    add_para(doc, "Date: 2026-05-19", italic=True, size=11)
    add_para(doc, "")

    add_heading(doc, "About this document", level=1)
    add_para(
        doc,
        "Two primary personas drive design decisions for DroneAid: one civilian impacted by conflict who requests relief supplies, and one relief coordinator who dispatches drones. Each persona is grounded in publicly reported scenarios from humanitarian-aid literature; names and identifying details are fictional.",
    )
    add_para(doc, "Personas inform the user journey map (04-journey-map.md), the function list (01-concepts.md), and acceptance criteria in the design spec.")

    # ── Persona 1 ──
    doc.add_page_break()
    add_heading(doc, "Persona 1 — Mali Suwan (End user)", level=1)
    add_photo_placeholder(doc, "Mali, 37")

    add_para(doc, "")
    add_kv_table(doc, [
        ("Age", "37"),
        ("Location", "Temporary shelter, displaced from home village 3 weeks ago"),
        ("Occupation", "Former small-shop owner; currently unemployed"),
        ("Household", "Two children (ages 6 and 9); elderly mother"),
        ("Languages", "Thai (native), basic English"),
        ("Devices", "Mid-range Android phone (3 years old), shared with mother; intermittent Wi-Fi at shelter, no mobile data plan"),
        ("Tech savvy", "Moderate — uses LINE daily, has used food-delivery apps, comfortable with QR codes"),
        ("Health/safety context", "Younger child has recurring fevers; aspirin and ORS in short supply at shelter; sometimes hears explosions in the distance"),
    ])

    add_heading(doc, "Goals", level=2)
    add_bullets(doc, [
        "Get critical supplies (medicine, ORS, baby food, water) for her children within hours, not days.",
        "Avoid leaving the shelter unnecessarily; ground travel feels unsafe.",
        "Know that her request was received and is being acted on — silence is worse than a delay.",
        "Track the drone so she can be ready when it lands and not miss the drop.",
    ])

    add_heading(doc, "Frustrations", level=2)
    add_bullets(doc, [
        "Forms that ask for ID numbers she doesn’t have ready (passport, household registry).",
        "Apps that work only over fast internet — hers is patchy.",
        "Not knowing why a request was rejected or delayed.",
        "Confusing maps that don’t show her exact shelter location.",
    ])

    add_heading(doc, "Behaviors", level=2)
    add_bullets(doc, [
        "Checks her phone every ~30 minutes when expecting something important.",
        "Reads notifications carefully; doesn’t dismiss them without acting.",
        "Cross-checks with neighbors before submitting forms — she’s cautious about giving information.",
        "Prefers Thai over English UI but can navigate either.",
    ])

    add_heading(doc, "Key quote", level=2)
    add_quote_box(
        doc,
        "I just need to know if the drone is really coming, and when. If I miss it, the children go another night without medicine.",
    )

    add_heading(doc, "Typical scenario", level=2)
    add_para(
        doc,
        "It is 8:00 p.m. Mali’s six-year-old has a fever. The shelter clinic ran out of paracetamol earlier today. Mali opens DroneAid, logs in with her national ID, picks one Medical Kit from the catalog, drops a pin on the shelter’s open field for delivery, and submits the request. She watches the queue page. Within minutes she gets a push: “Drone dispatched, ETA 18 minutes.” She opens the tracking page and watches the drone move across the map, battery dropping from 92% to 78% as it approaches. When it lands she retrieves the kit, taps Confirm, and goes back inside to give her daughter the medicine.",
    )

    add_heading(doc, "What DroneAid must do for Mali", level=2)
    add_bullets(doc, [
        "Register with national ID + password only — no documents to upload.",
        "Allow her to drop a pin on the map for the exact shelter spot.",
        "Show clear status: pending, approved, dispatched, arriving, delivered.",
        "Push a notification she will actually see; in-app inbox for replays.",
        "Live battery + ETA on tracking so she trusts the system.",
        "Explain failures: if a flight aborts due to storm, tell her, in plain language, that a new drone will be sent.",
    ])

    # ── Persona 2 ──
    doc.add_page_break()
    add_heading(doc, "Persona 2 — Naree Charoen (Admin / Relief Coordinator)", level=1)
    add_photo_placeholder(doc, "Naree, 29")

    add_para(doc, "")
    add_kv_table(doc, [
        ("Age", "29"),
        ("Location", "Forward operations base, ~12 km from the front line"),
        ("Occupation", "Relief operations coordinator, NGO; 4 years in field aid work"),
        ("Languages", "Thai, English, conversational Burmese"),
        ("Devices", "Issued ruggedized Android tablet; backup smartphone; reliable Starlink + 4G"),
        ("Tech savvy", "High — spreadsheets, dashboards, Slack, ATC-style coordination apps"),
        ("Team context", "Two-person dispatch desk; one drone-operator partner; on rotating 8-hour shifts"),
        ("Pressure", "Handles 30–80 requests per shift in busy periods; weather + fleet status change minute-to-minute"),
    ])

    add_heading(doc, "Goals", level=2)
    add_bullets(doc, [
        "Vet incoming requests quickly: who, what, where, payload feasible?",
        "Dispatch drones efficiently — maximize deliveries per shift with the smallest fleet.",
        "React to weather and failures in real time without losing track of requests in flight.",
        "Hand off cleanly to the next shift with an accurate fleet and queue state.",
    ])

    add_heading(doc, "Frustrations", level=2)
    add_bullets(doc, [
        "Hidden constraints that surface late — e.g. drone payload too small after picking.",
        "Tools that require many clicks to approve a request she’s already reviewed.",
        "Notifications she can’t silence per category (she wants failures loud, status pings quiet).",
        "Maps that hide flight paths or don’t show which drone is doing what.",
    ])

    add_heading(doc, "Behaviors", level=2)
    add_bullets(doc, [
        "Lives on the map view; switches to lists only when she has to.",
        "Sets weather state proactively when sensors at the warehouse change.",
        "Annotates rejections with reasons so the user sees something useful.",
        "Pre-positions drones during quiet windows to shorten future delivery times.",
    ])

    add_heading(doc, "Key quote", level=2)
    add_quote_box(
        doc,
        "I don’t care if it’s pretty. I care if I can approve, assign, and watch a drone in under ten seconds, and if I can see when one of my birds is in trouble.",
    )

    add_heading(doc, "Typical scenario", level=2)
    add_para(
        doc,
        "Naree opens DroneAid at the start of her shift. The Control screen shows three drones already in flight from the previous shift. She sets weather to “wind” because gusts have just picked up. A new request notification: Mali’s medical-kit request. She taps it, sees the user’s shelter location 4 km away, confirms stock and payload, and approves. The drone picker shows two eligible drones. She picks DRN-003 (full battery, closer base) and assigns. Twelve minutes later the system flags DRN-005 (another flight) for low-battery abort. Naree gets the alert, taps Reassign, dispatches DRN-007 in its place, and notes the rejection reason on the original flight log.",
    )

    add_heading(doc, "What DroneAid must do for Naree", level=2)
    add_bullets(doc, [
        "Surface new requests with a push she can act on with one tap.",
        "Filter drone picker to only eligible drones (idle, payload fits, in range).",
        "Make approve + assign two taps total when conditions are normal.",
        "Provide a live map showing every active drone, with battery + state at a glance.",
        "Alert clearly on flight failures and offer Reassign as the primary action.",
        "Let her set global weather state quickly — it should ripple through the sim instantly.",
    ])

    # ── Persona comparison table ──
    doc.add_page_break()
    add_heading(doc, "Comparison — design implications", level=1)
    add_para(doc, "Side-by-side of the two personas' needs and how DroneAid satisfies each.")
    add_para(doc, "")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(["Dimension", "Mali (User)", "Naree (Admin)"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(11)
        shade_cell(hdr[i], "1F4E78")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_border(hdr[i])

    rows = [
        ("Primary screen", "Tracking page — live drone + battery + ETA", "Control map + Request Manage"),
        ("Critical task", "Submit request and confirm receipt", "Approve, pick drone, react to failures"),
        ("Tap budget", "Many small interactions; tolerates wait time", "Optimize for taps-per-action; minimize friction"),
        ("Notifications", "Loud, all categories — she’s expecting something", "Category-tunable — failures loud, status quiet"),
        ("Map mode", "One drone, one flight", "All drones, all flights"),
        ("Failure tolerance", "Needs explanation + reassurance", "Needs action + recovery path"),
        ("Connectivity", "Patchy Wi-Fi, no data plan", "Reliable, redundant network"),
        ("Language", "Thai primary; English fallback", "Thai or English; switches by team"),
    ]
    for k, v, w in rows:
        row = table.add_row().cells
        for i, txt in enumerate((k, v, w)):
            row[i].text = ""
            p = row[i].paragraphs[0]
            r = p.add_run(txt)
            r.font.name = "Arial"
            r.font.size = Pt(10.5)
            set_cell_border(row[i])
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for row in table.rows:
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(2.6)
        row.cells[2].width = Inches(2.6)

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    build()
